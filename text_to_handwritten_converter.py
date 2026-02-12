import customtkinter as ctk
from tkinter import filedialog, messagebox
import os
import random
import datetime
import threading
import subprocess
import sys
from PIL import Image, ImageDraw, ImageFont
import docx
import pypdf

# --- 1. SYSTEM CONFIGURATION ---
ctk.set_appearance_mode("Dark") 
ctk.set_default_color_theme("dark-blue")

# --- 2. ASSET MANAGEMENT ---
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# --- FONTS LIST ---
AVAILABLE_FONTS = {
    "David Reid (Standard)": resource_path("fonts/QEDavidReid.ttf"),
    "David Reid (Caps)": resource_path("fonts/QEDavidReidCAP.ttf"),
    "Antony Lark": resource_path("fonts/QEAntonyLark.ttf"),
    "Braden Hill": resource_path("fonts/QEBradenHill.ttf"),
    "Caroline Mutiboko": resource_path("fonts/QECarolineMutiboko.ttf"),
    "Hughes Script": resource_path("fonts/QEGHHughes.ttf"),
    "Herbert Cooper": resource_path("fonts/QEHerbertCooper.ttf"),
    "Jeff Dungan": resource_path("fonts/QEJeffDungan.ttf"),
    "Scott Williams": resource_path("fonts/QEScottWilliams.ttf")
}

# Backgrounds
BUILTIN_BGS = {
    "Lined Paper": resource_path("backgrounds/img1.jpg")
}

# --- 3. GENERATION ENGINE ---
def generate_pdf(text, font_paths, bg_image_path, font_size, word_spacing, letter_spacing, font_color, line_variation):
    try:
        hw_fonts = []
        for p in font_paths:
            if os.path.exists(p):
                hw_fonts.append(ImageFont.truetype(p, size=font_size))
        
        if not hw_fonts:
            return "Error: No valid fonts found!"

        # Character Maps
        SUPERSCRIPTS = {
            '²':'2', '³':'3', '¹':'1', '⁰':'0', '⁴':'4', '⁵':'5', 
            '⁶':'6', '⁷':'7', '⁸':'8', '⁹':'9', '⁺':'+', '⁻':'-', 
            '⁼':'=', '⁽':'(', '⁾':')', 'ˣ':'x', 'ⁿ':'n'
        }
        SUBSCRIPTS = {
            '₀':'0', '₁':'1', '₂':'2', '₃':'3', '₄':'4', '₅':'5', 
            '₆':'6', '₇':'7', '₈':'8', '₉':'9', '₊':'+', '₋':'-', 
            '₌':'=', '₍':'(', '₎':')', 'ₓ':'x'
        }

        def get_new_page():
            if not bg_image_path or not os.path.exists(bg_image_path):
                img = Image.new("RGB", (2480, 3508), "white")
            else:
                img = Image.open(bg_image_path).convert("RGB")
            draw = ImageDraw.Draw(img)
            return img, draw

        # Layout Setup
        current_page_img, current_draw = get_new_page()
        page_width, page_height = current_page_img.size
        
        MARGIN_LEFT = int(page_width * 0.1)
        MARGIN_RIGHT = int(page_width * 0.1)
        MARGIN_TOP = int(page_height * 0.1)
        MARGIN_BOTTOM = int(page_height * 0.1)
        
        LINE_HEIGHT = int(font_size * 1.5)
        cursor_x, cursor_y = MARGIN_LEFT, MARGIN_TOP
        pages_list = []

        source_lines = text.split('\n')

        for line in source_lines:
            # Smart Tab/Indent Handling
            leading_spaces = len(line) - len(line.lstrip(' '))
            cursor_x = MARGIN_LEFT + (leading_spaces * (font_size // 3))

            if not line.strip():
                cursor_y += LINE_HEIGHT
                if cursor_y + LINE_HEIGHT > page_height - MARGIN_BOTTOM:
                    pages_list.append(current_page_img)
                    current_page_img, current_draw = get_new_page()
                    cursor_x, cursor_y = MARGIN_LEFT, MARGIN_TOP
                continue

            words = line.split(' ')
            for word in words:
                if not word: continue
                
                test_font = hw_fonts[0]
                word_width = test_font.getlength(word)

                # Add letter spacing estimation to word wrap calc
                total_word_width = word_width + (len(word) * letter_spacing)

                if cursor_x + total_word_width > page_width - MARGIN_RIGHT:
                    cursor_x = MARGIN_LEFT
                    cursor_y += LINE_HEIGHT

                if cursor_y + LINE_HEIGHT > page_height - MARGIN_BOTTOM:
                    pages_list.append(current_page_img)
                    current_page_img, current_draw = get_new_page()
                    cursor_x, cursor_y = MARGIN_LEFT, MARGIN_TOP

                for char in word:
                    jitter = random.randint(-line_variation, line_variation) if line_variation > 0 else 0

                    # CASE A: Superscript
                    if char in SUPERSCRIPTS:
                        normal_char = SUPERSCRIPTS[char]
                        base_font = random.choice(hw_fonts)
                        tiny_font = ImageFont.truetype(base_font.path, size=int(font_size * 0.6))
                        
                        lift_amount = int(font_size * 0.25) # Fixed Height
                        
                        current_draw.text((cursor_x, cursor_y - lift_amount + jitter), normal_char, font=tiny_font, fill=font_color)
                        cursor_x += tiny_font.getlength(normal_char) + 4 + letter_spacing

                    # CASE B: Subscript
                    elif char in SUBSCRIPTS:
                        normal_char = SUBSCRIPTS[char]
                        base_font = random.choice(hw_fonts)
                        tiny_font = ImageFont.truetype(base_font.path, size=int(font_size * 0.6))
                        
                        drop_amount = int(font_size * 0.5) 
                        
                        current_draw.text((cursor_x, cursor_y + drop_amount + jitter), normal_char, font=tiny_font, fill=font_color)
                        cursor_x += tiny_font.getlength(normal_char) + 4 + letter_spacing

                    # CASE C: Normal Character
                    else:
                        selected_font = random.choice(hw_fonts)
                        stamp_size = int(font_size * 4)
                        txt_img = Image.new('RGBA', (stamp_size, stamp_size), (255, 255, 255, 0))
                        txt_draw = ImageDraw.Draw(txt_img)
                        txt_draw.text((stamp_size/2, stamp_size/2), char, font=selected_font, fill=font_color+(255,), anchor="ms")
                        
                        angle = random.randint(-2, 2)
                        rotated = txt_img.rotate(angle, expand=1, resample=Image.BICUBIC)
                        
                        paste_y = int((cursor_y + (font_size * 1.1)) - (stamp_size/2)) + jitter
                        paste_x = int(cursor_x - (stamp_size/2))
                        
                        current_page_img.paste(rotated, (paste_x, paste_y), rotated)
                        
                        # Add Character Width + Letter Spacing + Random Micro-Gap
                        cursor_x += selected_font.getlength(char) + letter_spacing + random.randint(0, 1)

                cursor_x += word_spacing + random.randint(-2, 2)

            cursor_x = MARGIN_LEFT
            cursor_y += LINE_HEIGHT
            
            if cursor_y + LINE_HEIGHT > page_height - MARGIN_BOTTOM:
                pages_list.append(current_page_img)
                current_page_img, current_draw = get_new_page()
                cursor_x, cursor_y = MARGIN_LEFT, MARGIN_TOP

        pages_list.append(current_page_img)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        out_name = f"Handwritten_{timestamp}.pdf"
        pages_list[0].save(out_name, "PDF", resolution=100.0, save_all=True, append_images=pages_list[1:])
        return out_name
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return str(e)

# --- 4. APP UI ---
class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Handwrite.ai - Assignment Converter")
        self.geometry("950x880") # Increased height slightly
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        self.font_vars = {}
        self.selected_bg_path = None

        # HEADER
        self.header = ctk.CTkFrame(self, corner_radius=0, fg_color="#1a1a1a")
        self.header.grid(row=0, column=0, sticky="ew")
        ctk.CTkLabel(self.header, text="Handwrite.ai", font=("Roboto Medium", 24)).pack(pady=15)

        # SCROLL AREA
        self.main_scroll = ctk.CTkScrollableFrame(self, fg_color="transparent")
        self.main_scroll.grid(row=1, column=0, sticky="nsew", padx=20, pady=20)
        self.main_scroll.grid_columnconfigure(0, weight=1)

        # CARD 1: INPUT
        self.card_input = ctk.CTkFrame(self.main_scroll)
        self.card_input.grid(row=0, column=0, sticky="ew", pady=(0, 15))
        self.card_input.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(self.card_input, text="1. Input Text", font=("Arial", 16, "bold")).grid(row=0, column=0, sticky="w", padx=15, pady=10)
        self.btn_import = ctk.CTkButton(self.card_input, text="Import from Word/PDF", command=self.import_file, fg_color="#2b2b2b", border_width=1)
        self.btn_import.grid(row=0, column=1, padx=15, pady=10)

        self.textbox = ctk.CTkTextbox(self.card_input, height=180, font=("Consolas", 12))
        self.textbox.grid(row=1, column=0, columnspan=2, padx=15, pady=(0, 15), sticky="ew")
        self.textbox.insert("0.0", "Paste text here or import a file...")

        # CARD 2: SETTINGS
        self.card_config = ctk.CTkFrame(self.main_scroll)
        self.card_config.grid(row=1, column=0, sticky="ew", pady=(0, 15))
        self.card_config.grid_columnconfigure(0, weight=1)
        self.card_config.grid_columnconfigure(1, weight=1)

        # Fonts
        ctk.CTkLabel(self.card_config, text="2. Handwriting Style", font=("Arial", 16, "bold")).grid(row=0, column=0, sticky="w", padx=15, pady=10)
        self.frame_fonts = ctk.CTkScrollableFrame(self.card_config, height=220) # Increased height for more fonts
        self.frame_fonts.grid(row=1, column=0, padx=15, pady=10, sticky="ew")
        
        for name, path in AVAILABLE_FONTS.items():
            var = ctk.BooleanVar(value=True if "David" in name else False)
            chk = ctk.CTkCheckBox(self.frame_fonts, text=name, variable=var)
            chk.pack(anchor="w", pady=2)
            self.font_vars[path] = var

        # Page Settings
        ctk.CTkLabel(self.card_config, text="3. Page Settings", font=("Arial", 16, "bold")).grid(row=0, column=1, sticky="w", padx=15, pady=10)
        self.frame_sets = ctk.CTkFrame(self.card_config, fg_color="transparent")
        self.frame_sets.grid(row=1, column=1, padx=15, pady=10, sticky="new")

        self.bg_options = list(BUILTIN_BGS.keys()) + ["Import Custom Image..."]
        self.combo_bg = ctk.CTkComboBox(self.frame_sets, values=self.bg_options, command=self.handle_bg_selection)
        self.combo_bg.pack(fill="x", pady=(0, 10))
        self.combo_bg.set("Lined Paper")
        self.selected_bg_path = BUILTIN_BGS["Lined Paper"]

        ctk.CTkLabel(self.frame_sets, text="Font Size").pack(anchor="w")
        self.slider_size = ctk.CTkSlider(self.frame_sets, from_=30, to=100)
        self.slider_size.set(50)
        self.slider_size.pack(fill="x", pady=5)

        # --- NEW: LETTER SPACING ---
        ctk.CTkLabel(self.frame_sets, text="Letter Spacing (Character Gap)").pack(anchor="w")
        self.slider_letter_spacing = ctk.CTkSlider(self.frame_sets, from_=0, to=10)
        self.slider_letter_spacing.set(0) # Default 0
        self.slider_letter_spacing.pack(fill="x", pady=5)
        
        ctk.CTkLabel(self.frame_sets, text="Word Spacing").pack(anchor="w")
        self.slider_spacing = ctk.CTkSlider(self.frame_sets, from_=10, to=100)
        self.slider_spacing.set(40)
        self.slider_spacing.pack(fill="x", pady=5)
        
        ctk.CTkLabel(self.frame_sets, text="Line Variation (Messiness)").pack(anchor="w")
        self.slider_variation = ctk.CTkSlider(self.frame_sets, from_=0, to=15)
        self.slider_variation.set(5) 
        self.slider_variation.pack(fill="x", pady=5)

        ctk.CTkLabel(self.frame_sets, text="Ink Color").pack(anchor="w")
        self.combo_color = ctk.CTkComboBox(self.frame_sets, values=["Blue", "Black", "Red"])
        self.combo_color.pack(fill="x", pady=10)

        # FOOTER
        self.footer = ctk.CTkFrame(self, fg_color="transparent")
        self.footer.grid(row=2, column=0, sticky="ew", padx=20, pady=10)
        
        self.progress = ctk.CTkProgressBar(self.footer, mode="indeterminate")
        self.progress.pack(fill="x", pady=5)
        self.progress.set(0)

        self.btn_run = ctk.CTkButton(self.footer, text="GENERATE PDF DOCUMENT", height=50, font=("Arial", 16, "bold"), command=self.start_generation_thread)
        self.btn_run.pack(fill="x")

    def handle_bg_selection(self, choice):
        if choice == "Import Custom Image...":
            path = filedialog.askopenfilename(filetypes=[("Image Files", "*.jpg *.png")])
            if path:
                self.selected_bg_path = path
            else:
                self.combo_bg.set("Lined Paper")
                self.selected_bg_path = BUILTIN_BGS["Lined Paper"]
        elif choice in BUILTIN_BGS:
            self.selected_bg_path = BUILTIN_BGS[choice]

    def import_file(self):
        path = filedialog.askopenfilename(filetypes=[("Documents", "*.docx *.pdf")])
        if not path: return
        try:
            text = ""
            if path.endswith(".docx"):
                doc = docx.Document(path)
                full_text_list = []
                SUPERSCRIPT_MAP = {'0':'⁰', '1':'¹', '2':'²', '3':'³', '4':'⁴', '5':'⁵', '6':'⁶', '7':'⁷', '8':'⁸', '9':'⁹', '+':'⁺', '-':'⁻', '=':'⁼', '(':'⁽', ')':'⁾', 'x':'ˣ', 'n':'ⁿ'}
                SUBSCRIPT_MAP = {'0':'₀', '1':'₁', '2':'₂', '3':'₃', '4':'₄', '5':'₅', '6':'₆', '7':'₇', '8':'₈', '9':'₉', '+':'₊', '-':'₋', '=':'₌', '(':'₍', ')':'₎', 'x':'ₓ'}

                para_iter = iter(doc.paragraphs)
                table_iter = iter(doc.tables)

                for element in doc.element.body.iterchildren():
                    if element.tag.endswith('p'):
                        try:
                            para = next(para_iter)
                            indent_spaces = ""
                            try:
                                if para.paragraph_format.left_indent:
                                    inches = para.paragraph_format.left_indent.inches
                                    if inches: indent_spaces += " " * int(inches * 12)
                                if para.paragraph_format.first_line_indent:
                                    inches = para.paragraph_format.first_line_indent.inches
                                    if inches and inches > 0: indent_spaces += " " * int(inches * 12)
                            except: pass
                                
                            para_string = indent_spaces
                            for run in para.runs:
                                run_text = run.text.replace("\t", "    ")
                                if run.font.superscript:
                                    para_string += "".join([SUPERSCRIPT_MAP.get(c, c) for c in run_text])
                                elif run.font.subscript:
                                    para_string += "".join([SUBSCRIPT_MAP.get(c, c) for c in run_text])
                                else:
                                    para_string += run_text
                            full_text_list.append(para_string)
                        except StopIteration: pass
                    elif element.tag.endswith('tbl'):
                        try:
                            table = next(table_iter)
                            full_text_list.append("\n" + " "*5 + "-"*30) 
                            for row in table.rows:
                                row_text = "   |   ".join([cell.text.strip() for cell in row.cells])
                                full_text_list.append("   " + row_text)
                                full_text_list.append("   " + "-" * len(row_text))
                            full_text_list.append(" "*5 + "-"*30 + "\n")
                        except StopIteration: pass
                text = "\n".join(full_text_list)
            elif path.endswith(".pdf"):
                reader = pypdf.PdfReader(path)
                text = "\n".join([p.extract_text() for p in reader.pages])
            self.textbox.delete("0.0", "end")
            self.textbox.insert("0.0", text)
        except Exception as e:
            messagebox.showerror("Error", f"Import failed: {e}")

    def start_generation_thread(self):
        text = self.textbox.get("0.0", "end").strip()
        if len(text) < 5:
            messagebox.showwarning("Warning", "Text is too short!")
            return
        selected_fonts = [p for p, v in self.font_vars.items() if v.get()]
        if not selected_fonts:
            messagebox.showwarning("Warning", "Select at least one font!")
            return

        self.btn_run.configure(state="disabled", text="Generating... (This may take a minute)")
        self.progress.start()
        threading.Thread(target=self.run_generation_logic, args=(text, selected_fonts), daemon=True).start()

    def run_generation_logic(self, text, fonts):
        size = int(self.slider_size.get())
        word_spacing = int(self.slider_spacing.get())
        letter_spacing = int(self.slider_letter_spacing.get()) # GET NEW SLIDER
        variation = int(self.slider_variation.get())
        color_map = {"Blue": (0, 0, 100), "Black": (20, 20, 20), "Red": (150, 0, 0)}
        color = color_map[self.combo_color.get()]

        result = generate_pdf(text, fonts, self.selected_bg_path, size, word_spacing, letter_spacing, color, variation)
        self.after(0, lambda: self.finish_generation(result))

    def finish_generation(self, result):
        self.progress.stop()
        self.progress.set(0)
        self.btn_run.configure(state="normal", text="GENERATE PDF DOCUMENT")
        if result.endswith(".pdf"):
            if messagebox.askyesno("Success", f"PDF Generated!\n{result}\n\nOpen containing folder?"):
                subprocess.Popen(f'explorer /select,"{os.path.abspath(result)}"')
        else:
            messagebox.showerror("Error", result)

if __name__ == "__main__":
    app = App()

    app.mainloop()
